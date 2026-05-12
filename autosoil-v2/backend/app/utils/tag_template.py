# app/utils/tag_template.py
import os
from docx import Document

def prepare_template(input_path: str, output_path: str):
    """
    A utility script to inject docxtpl Jinja2 tags into an existing Geolog Word template.
    This replaces placeholder text with {{ variable_name }} so the swarm can auto-fill it.
    """
    if not os.path.exists(input_path):
        print(f"Error: Could not find template at {input_path}")
        return

    doc = Document(input_path)
    
    # Example logic: replace literal strings or add a table block at the end.
    # For a real template, you would replace specific bookmarks or text patterns.
    # Here we just append the necessary jinja loop for layers to prove it works.
    
    doc.add_heading('AutoSoil Logged Data (Auto-Generated)', level=1)
    doc.add_paragraph('Project: {{ project_name }}')
    doc.add_paragraph('Borehole: {{ borehole_id }}')
    doc.add_paragraph('Date: {{ date_logged }}')
    
    doc.add_paragraph('{% tr for layer in soil_layers %}')
    doc.add_paragraph('Depth: {{ layer.depth_from }}m to {{ layer.depth_to }}m')
    doc.add_paragraph('USCS: {{ layer.uscs_code }}')
    doc.add_paragraph('Description: {{ layer.description }}')
    doc.add_paragraph('{% tr endfor %}')
    
    doc.save(output_path)
    print(f"Successfully created tagged template at: {output_path}")

if __name__ == "__main__":
    source = r"C:\Users\pored\Downloads\Project Geologs\Templates\Geotechnical Investigation Master Template.docx"
    dest = r"C:\Users\pored\Downloads\Project Geologs\Templates\Geotechnical Investigation Master Template_TAGGED.docx"
    prepare_template(source, dest)
